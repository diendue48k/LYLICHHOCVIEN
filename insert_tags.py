import docx
import re
import os

def replace_text_in_doc(doc_path, output_path, is_ncs):
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Cannot open {doc_path}: {e}")
        return

    # Basic replacements
    replacements = {
        "Họ và tên:": "Họ và tên: {fullName}",
        "Ngày sinh:": "Ngày sinh: {birthDate_formatted}",
        "Nam / Nữ": "Nam: {gender_nam}  Nữ: {gender_nu}",
        "Nam/Nữ": "Nam: {gender_nam}  Nữ: {gender_nu}",
        "Nơi sinh:": "Nơi sinh: {birthPlace}",
        "Quốc tịch:": "Quốc tịch: {nationality}",
        "Dân tộc:": "Dân tộc: {ethnic}",
        "Tôn giáo:": "Tôn giáo: {religion}",
        "Số thẻ Căn cước công dân (hoặc CMND):": "Số CCCD: {cccd}",
        "Số thẻ Căn cước công dân:": "Số CCCD: {cccd}",
        "Số CMND:": "Số CMND: {cmnd}",
        "Cấp ngày:": "Cấp ngày: {cccdIssuedDate_formatted}",
        "Nơi cấp:": "Nơi cấp: {cccdIssuedPlace}",
        "Có /Không": "Có: {disability_co}   Không: {disability_khong}",
        "Loại khuyết tật:": "Loại khuyết tật: {disabilityType}",
        "Quê quán:": "Quê quán: {hometown}",
        "Số sổ bảo hiểm xã hội:": "Số sổ bảo hiểm xã hội: {socialInsuranceId}",
        "Ngày vào Đoàn TNCS HCM:": "Ngày vào Đoàn TNCS HCM: {unionJoinDate_formatted}",
        "Ngày vào Đảng CSVN:": "Ngày vào Đảng CSVN: {partyJoinDate_formatted}",
        "Ngày chính thức:": "Ngày chính thức: {officialPartyJoinDate_formatted}",
        "Cơ quan công tác hiện nay:": "Cơ quan công tác hiện nay: {workPlace}",
        "Số ĐT cá nhân (cần khai chính xác):": "Số điện thoại: {phone}",
        "Số điện thoại nghiên cứu sinh:": "Số điện thoại nghiên cứu sinh: {phone}",
        "Email:": "Email: {email}",
        "Họ tên người thân khi cần liên hệ:": "Họ tên người thân khi cần liên hệ: {emergencyContactName}",
        "Số ĐT:": "Số ĐT: {emergencyContactPhone}",
        "Số điện thoại của người thân:": "Số điện thoại của người thân: {emergencyContactPhone}",
        "Hộ khẩu thường trú:": "Hộ khẩu thường trú: Số nhà: {perm_house} Đường: {perm_street} Tổ/Thôn: {perm_hamlet} Phường/Xã: {perm_ward} Quận/Huyện: {perm_district} Tỉnh/TP: {perm_city}",
        "Địa chỉ ở hiện nay:": "Địa chỉ ở hiện nay: Số nhà: {curr_house} Đường: {curr_street} Tổ/Thôn: {curr_hamlet} Phường/Xã: {curr_ward} Quận/Huyện: {curr_district} Tỉnh/TP: {curr_city}",
        
        "Trường cấp bằng:": "Trường cấp bằng: {uni_school}",
        "Hệ đào tạo:": "Hệ đào tạo: {uni_system}",
        "Số hiệu bằng:": "Số hiệu bằng: {uni_number}",
        "Số vào sổ cấp bằng:": "Số vào sổ cấp bằng: {uni_book}",
        "Ngày ký bằng:": "Ngày ký bằng: {uni_date_formatted}",
        "Năm tốt nghiệp:": "Năm tốt nghiệp: {uni_year}",
        "Ngành:": "Ngành: {uni_major}",
        "Chuyên ngành:": "Chuyên ngành: {uni_spec}",
        "Học hàm-Học vị/ Họ tên người ký bằng:": "Người ký bằng: {uni_signer}",
        "Học hàm-Học vị/Họ tên người ký bằng:": "Người ký bằng: {uni_signer}",
        "Là bằng liên thông:": "Là bằng liên thông: Có {uni_transfer_co} Không {uni_transfer_khong}"
    }

    if is_ncs:
        replacements.update({
            "Trường cấp bằng Thạc sĩ:": "Trường cấp bằng: {master_school}",
            "Hệ đào tạo Thạc sĩ:": "Hệ đào tạo: {master_system}",
            "Số hiệu bằng Thạc sĩ:": "Số hiệu bằng: {master_number}",
            "Số vào sổ cấp bằng Thạc sĩ:": "Số vào sổ cấp bằng: {master_book}",
            "Ngày ký bằng Thạc sĩ:": "Ngày ký bằng: {master_date_formatted}",
            "Năm tốt nghiệp Thạc sĩ:": "Năm tốt nghiệp: {master_year}",
            "Ngành Thạc sĩ:": "Ngành: {master_major}",
            "Chuyên ngành Thạc sĩ:": "Chuyên ngành: {master_spec}",
        })

    def replace_in_p(p):
        text = p.text
        if not text.strip():
            return
        
        # Simple rule: if we see a label, we append the tag at the end of the paragraph.
        # But wait, it's better to clear the runs and just set the text so we don't duplicate.
        for key, val in replacements.items():
            if key in text and "{" not in text:
                text = text.replace(key, val)
                # Remove dotted lines
                text = re.sub(r'\.{3,}', '', text)
                p.text = text
                break # only replace one per paragraph to avoid mess, or continue?

    for p in doc.paragraphs:
        replace_in_p(p)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_p(p)

    doc.save(output_path)

replace_text_in_doc('1. Phiếu Lý lịch học viên.docx', 'public/1. Phiếu Lý lịch học viên.docx', False)
replace_text_in_doc('4. Phiếu thông tin Lý lịch NCS.docx', 'public/4. Phiếu thông tin Lý lịch NCS.docx', True)
