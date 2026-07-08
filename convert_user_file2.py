import docx
import re

def convert_user_file(doc_path, output_path):
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Cannot open {doc_path}: {e}")
        return

    # Map of their exact typed texts to the tags
    replacements = {
        "TRẦN VĂN NAM": "{fullName}",
        "20/08/1959": "{birthDate_formatted}",
        "☑ Nam   Nữ": "{gender_nam} Nam   {gender_nu} Nữ",
        "☑ Nam": "{gender_nam} Nam",
        "Nữ": "{gender_nu} Nữ",
        "Việt Nam": "{nationality}",
        "Không": "{religion}",
        "048187007211": "{cccd}",
        "28/09/2021": "{cccdIssuedDate_formatted}",
        "Cục Cảnh sát QLHC về TTXH": "{cccdIssuedPlace}",
        "Có   ☑  Không": "Có {disability_co}   Không {disability_khong}",
        "Có   ☑": "Có {disability_co}",
        "Phường Ngũ Hành Sơn, thành phố Đà Nẵng": "{hometown}",
        "4810000577": "{socialInsuranceId}",
        "26/03/2006": "{unionJoinDate_formatted}",
        "19/05/2015": "{partyJoinDate_formatted}",
        "19/05/2016": "{officialPartyJoinDate_formatted}",
        "Trường Đại học Kinh tế - Đại học Đà Nẵng": "{workPlace}",
        "0969754149": "{phone}",
        "namtv@due.udn.vn": "{email}",
        "Trần Văn Tính": "{emergencyContactName}",
        "0969764149": "{emergencyContactPhone}",
        "71": "{perm_house}",
        "Tổ 54": "{perm_hamlet}",
        
        # Bang Tot Nghiep
        "Chính quy": "{uni_system}",
        "001453": "{uni_number}",
        "CQ09/0066": "{uni_book}",
        "29/06/2006": "{uni_date_formatted}",
        "2006": "{uni_year}",
        "GS.TS Lê Văn Huy": "{uni_signer}",
        "Có      ☑ Không": "Có {uni_transfer_co}      Không {uni_transfer_khong}",
        
        # Gia Dinh
        "TRẦN VĂN A": "{father_name}",
        "20/9/1932": "{father_yob}",
        "0969555147": "{father_phone}",
        "Hưu trí": "{father_job}",
        "71 Ngũ Hành Sơn, Phường Ngũ Hành Sơn, thành phố Đà Nẵng": "{father_address}",
        "☑ Chưa tốt nghiệp PTTH": "☑ {father_edu_chuatn}",
        "TRẦN THỊ C": "{mother_name}",
        "0969555167": "{mother_phone}",
        "TRẦN THỊ H": "{spouse_name}",
        "20/9/1960": "{spouse_yob}",
        "0969555117": "{spouse_phone}",
        "☑ Đại học": "☑ {spouse_edu_dh}",
    }

    def process_p(p):
        if not p.text.strip():
            return
            
        # Merge runs
        text = p.text
        if len(p.runs) > 0:
            p.runs[0].text = text
            for i in range(1, len(p.runs)):
                p.runs[i].text = ""

        run = p.runs[0]
        
        # Strict word boundary for 'Kinh' and 'Kế toán'
        if re.search(r'\bKinh\b', run.text):
            if "Dân tộc:" in text or "Kinh" == text.strip():
                run.text = re.sub(r'\bKinh\b', "{ethnic}", run.text)
                
        if re.search(r'\bKế toán\b', run.text):
            if "Ngành:" in text:
                run.text = re.sub(r'\bKế toán\b', "{uni_major}", run.text)
            if "Chuyên ngành:" in text:
                run.text = re.sub(r'\bKế toán\b', "{uni_spec}", run.text)
                
        # Fix the bad replacement made previously if it's already there (though we are re-processing the original)
        # Actually, let's fix if we are modifying their edited file
        if "{ethnic} tế phát triển" in run.text:
            run.text = run.text.replace("{ethnic} tế phát triển", "{uni_major}")
            
        if "K49.KPT.NCS5" in run.text:
            run.text = run.text.replace("K49.KPT.NCS5", "{className}")
            
        if "Nơi sinh: Đà Nẵng" in run.text:
            run.text = run.text.replace("Đà Nẵng", "{birthPlace}")
        elif "Tên đường: Ngũ Hành Sơn" in run.text:
            run.text = run.text.replace("Ngũ Hành Sơn", "{perm_street}")
        elif "Xã/Phường: Ngũ Hành Sơn" in run.text:
            run.text = run.text.replace("Ngũ Hành Sơn", "{perm_ward}")
        elif "Tỉnh/Thành phố: Đà Nẵng" in run.text:
            run.text = run.text.replace("Đà Nẵng", "{perm_city}")
            
        # Only replace 'Không' for religion
        if "Tôn giáo: Không" in run.text:
            run.text = run.text.replace("Tôn giáo: Không", "Tôn giáo: {religion}")
            
        for k, v in replacements.items():
            if k == "Không": continue # skip the global one
            if k in run.text:
                if k == "Trường Đại học Kinh tế - Đại học Đà Nẵng" and "Cơ quan công tác" not in text and "Trường cấp bằng" not in text:
                    continue
                run.text = run.text.replace(k, v)

    for p in doc.paragraphs:
        process_p(p)
        
    for table in doc.tables:
        if len(table.rows) > 1 and "Thời gian" in table.rows[0].cells[0].text:
            if "trường" in table.rows[0].cells[1].text.lower() or "cơ sở" in table.rows[0].cells[1].text.lower():
                table.rows[1].cells[0].text = "{#educationHistory}{timeRange}"
                table.rows[1].cells[1].text = "{schoolName}"
                table.rows[1].cells[2].text = "{major}"
                table.rows[1].cells[3].text = "{learningType}"
                table.rows[1].cells[4].text = "{degree}{/educationHistory}"
                # delete other rows
                for i in range(len(table.rows)-1, 1, -1):
                    table._element.remove(table.rows[i]._element)
            elif "đơn vị" in table.rows[0].cells[1].text.lower():
                table.rows[1].cells[0].text = "{#workHistory}{timeRange}"
                table.rows[1].cells[1].text = "{organization}"
                table.rows[1].cells[2].text = "{position}{/workHistory}"
                # delete other rows
                for i in range(len(table.rows)-1, 1, -1):
                    table._element.remove(table.rows[i]._element)
        else:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_p(p)
                        
    for i, p in enumerate(doc.paragraphs):
        if "VI. CÁC CÔNG TRÌNH KHOA HỌC" in p.text:
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = "{#scientificWorks}{.}{/scientificWorks}"
        if "VII. KHEN THƯỞNG, KỈ LUẬT" in p.text:
            pass
        if "Khen thưởng: " in p.text:
            doc.paragraphs[i].text = "Khen thưởng: {rewards}"
        if "Kỷ luật: " in p.text:
            doc.paragraphs[i].text = "Kỷ luật: {discipline}"

    doc.save(output_path)
    
convert_user_file('1. Phiếu Lý lịch học viên.docx', 'public/1_Template_Mailing_v3.docx')
