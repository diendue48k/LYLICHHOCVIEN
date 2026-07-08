import sys
from docx import Document
from docx.shared import Pt

def fix_fonts(docx_path):
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                                
        doc.save(docx_path)
        print(f"Fixed {docx_path}")
    except Exception as e:
        print(f"Error processing {docx_path}: {e}")

if __name__ == '__main__':
    fix_fonts(sys.argv[1])
