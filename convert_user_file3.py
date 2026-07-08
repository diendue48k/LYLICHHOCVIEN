import docx
import re

def convert_user_file(doc_path, output_path):
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Cannot open {doc_path}: {e}")
        return

    def process_p(p, state):
        if not p.text.strip():
            return
            
        # Merge runs
        text = p.text
        if len(p.runs) > 0:
            p.runs[0].text = text
            for i in range(1, len(p.runs)):
                p.runs[i].text = ""

        run = p.runs[0]
        
        # State transitions
        if "Họ và tên MẸ:" in run.text:
            state = 'MOTHER'
        elif "Họ và tên CHỒNG/VỢ:" in run.text:
            state = 'SPOUSE'
        elif "Họ và tên CHA:" in run.text:
            state = 'FATHER'
            
        # Global replacements
        run.text = run.text.replace("TRẦN VĂN NAM", "{fullName}")
        run.text = run.text.replace("20/08/1959", "{birthDate_formatted}")
        run.text = run.text.replace("☑ Nam   Nữ", "{gender_nam} Nam   {gender_nu} Nữ")
        run.text = run.text.replace("☑ Nam", "{gender_nam} Nam")
        run.text = run.text.replace("Nữ", "{gender_nu} Nữ")
        run.text = run.text.replace("Việt Nam", "{nationality}")
        run.text = run.text.replace("048187007211", "{cccd}")
        run.text = run.text.replace("28/09/2021", "{cccdIssuedDate_formatted}")
        run.text = run.text.replace("Cục Cảnh sát QLHC về TTXH", "{cccdIssuedPlace}")
        
        if "Có   ☑  Không" in run.text:
            run.text = run.text.replace("Có   ☑  Không", "Có {disability_co}   Không {disability_khong}")
        elif "Có   ☑" in run.text:
            run.text = run.text.replace("Có   ☑", "Có {disability_co}")
            
        run.text = run.text.replace("4810000577", "{socialInsuranceId}")
        run.text = run.text.replace("26/03/2006", "{unionJoinDate_formatted}")
        run.text = run.text.replace("19/05/2015", "{partyJoinDate_formatted}")
        run.text = run.text.replace("19/05/2016", "{officialPartyJoinDate_formatted}")
        run.text = run.text.replace("0969754149", "{phone}")
        run.text = run.text.replace("namtv@due.udn.vn", "{email}")
        run.text = run.text.replace("Trần Văn Tính", "{emergencyContactName}")
        run.text = run.text.replace("0969764149", "{emergencyContactPhone}")
        
        # Bang Tot Nghiep
        run.text = run.text.replace("Chính quy", "{uni_system}")
        run.text = run.text.replace("001453", "{uni_number}")
        run.text = run.text.replace("CQ09/0066", "{uni_book}")
        run.text = run.text.replace("29/06/2006", "{uni_date_formatted}")
        run.text = run.text.replace("2006", "{uni_year}")
        run.text = run.text.replace("GS.TS Lê Văn Huy", "{uni_signer}")
        if "Có      ☑ Không" in run.text:
            run.text = run.text.replace("Có      ☑ Không", "Có {uni_transfer_co}      Không {uni_transfer_khong}")

        if re.search(r'\bKinh\b', run.text) and "Dân tộc:" in text:
            run.text = re.sub(r'\bKinh\b', "{ethnic}", run.text)
            
        if re.search(r'\bKế toán\b', run.text):
            if "Ngành:" in text:
                run.text = re.sub(r'\bKế toán\b', "{uni_major}", run.text)
            if "Chuyên ngành:" in text:
                run.text = re.sub(r'\bKế toán\b', "{uni_spec}", run.text)
                
        if "Nơi sinh: Đà Nẵng" in run.text:
            run.text = run.text.replace("Đà Nẵng", "{birthPlace}")
            
        # Fix addresses
        if "Quê quán: Phường Ngũ Hành Sơn, thành phố Đà Nẵng" in run.text:
            run.text = run.text.replace("Phường Ngũ Hành Sơn, thành phố Đà Nẵng", "{hometown}")
            
        if state == 'GENERAL':
            if "Tôn giáo: Không" in run.text:
                run.text = run.text.replace("Tôn giáo: Không", "Tôn giáo: {religion}")
            if "Tên đường: Ngũ Hành Sơn" in run.text:
                run.text = run.text.replace("Ngũ Hành Sơn", "{perm_street}")
            if "Xã/Phường: Ngũ Hành Sơn" in run.text:
                run.text = run.text.replace("Ngũ Hành Sơn", "{perm_ward}")
            if "Tỉnh/Thành phố: Đà Nẵng" in run.text:
                run.text = run.text.replace("Đà Nẵng", "{perm_city}")
            if "Số nhà: 71" in run.text:
                run.text = run.text.replace("71", "{perm_house}")
            if "Tổ/Thôn/Xóm: Tổ 54" in run.text:
                run.text = run.text.replace("Tổ 54", "{perm_hamlet}")
            if "Trường Đại học Kinh tế - Đại học Đà Nẵng" in run.text:
                if "Cơ quan công tác" in text:
                    run.text = run.text.replace("Trường Đại học Kinh tế - Đại học Đà Nẵng", "{workPlace}")
                elif "Trường cấp bằng" in text:
                    run.text = run.text.replace("Trường Đại học Kinh tế - Đại học Đà Nẵng", "{uni_school}")
                    
        # State-based Family replacements
        if state in ['FATHER', 'MOTHER', 'SPOUSE']:
            prefix = ""
            if state == 'FATHER': prefix = "father"
            elif state == 'MOTHER': prefix = "mother"
            elif state == 'SPOUSE': prefix = "spouse"
            
            # Name
            run.text = run.text.replace("TRẦN VĂN A", "{"+prefix+"_name}")
            run.text = run.text.replace("TRẦN THỊ C", "{"+prefix+"_name}")
            run.text = run.text.replace("TRẦN THỊ H", "{"+prefix+"_name}")
            # YOB
            run.text = run.text.replace("20/9/1932", "{"+prefix+"_yob}")
            run.text = run.text.replace("20/9/1960", "{"+prefix+"_yob}")
            # Phone
            run.text = run.text.replace("0969555147", "{"+prefix+"_phone}")
            run.text = run.text.replace("0969555167", "{"+prefix+"_phone}")
            run.text = run.text.replace("0969555117", "{"+prefix+"_phone}")
            # Job
            run.text = run.text.replace("Hưu trí", "{"+prefix+"_job}")
            # Address (the whole string)
            run.text = run.text.replace("71 Ngũ Hành Sơn, Phường Ngũ Hành Sơn, thành phố Đà Nẵng", "{"+prefix+"_address}")
            
            # Education checkboxes (recreate them accurately)
            # They wrote:
            # Sau đại học 	 Đại học           Cao đẳng        Trung cấp
            # Phổ thông trung học                       ☑ Chưa tốt nghiệp PTTH
            if "Sau đại học" in run.text and "Đại học" in run.text and "Cao đẳng" in run.text:
                run.text = run.text.replace("Sau đại học", "{"+prefix+"_edu_sdh} Sau đại học")
                run.text = run.text.replace("☑ Đại học", "{"+prefix+"_edu_dh} Đại học")
                if "☑ Đại học" not in text: # If not checked
                    run.text = run.text.replace("Đại học", "{"+prefix+"_edu_dh} Đại học")
                run.text = run.text.replace("Cao đẳng", "{"+prefix+"_edu_cd} Cao đẳng")
                run.text = run.text.replace("Trung cấp", "{"+prefix+"_edu_tc} Trung cấp")
                
            if "Phổ thông trung học" in run.text and "Chưa tốt nghiệp PTTH" in run.text:
                run.text = run.text.replace("Phổ thông trung học", "{"+prefix+"_edu_ptth} Phổ thông trung học")
                run.text = run.text.replace("☑ Chưa tốt nghiệp PTTH", "{"+prefix+"_edu_chuatn} Chưa tốt nghiệp PTTH")
                if "☑ Chưa tốt nghiệp PTTH" not in text:
                    run.text = run.text.replace("Chưa tốt nghiệp PTTH", "{"+prefix+"_edu_chuatn} Chưa tốt nghiệp PTTH")
                    
        return state

    state = 'GENERAL'
    for p in doc.paragraphs:
        state = process_p(p, state)
        
    for table in doc.tables:
        if len(table.rows) > 1 and "Thời gian" in table.rows[0].cells[0].text:
            if "trường" in table.rows[0].cells[1].text.lower() or "cơ sở" in table.rows[0].cells[1].text.lower():
                table.rows[1].cells[0].text = "{#educationHistory}{timeRange}"
                table.rows[1].cells[1].text = "{schoolName}"
                table.rows[1].cells[2].text = "{major}"
                table.rows[1].cells[3].text = "{learningType}"
                table.rows[1].cells[4].text = "{degree}{/educationHistory}"
                # delete other rows
                for i in range(len(table.rows)-1, 1, -1):
                    table._element.remove(table.rows[i]._element)
            elif "đơn vị" in table.rows[0].cells[1].text.lower():
                table.rows[1].cells[0].text = "{#workHistory}{timeRange}"
                table.rows[1].cells[1].text = "{organization}"
                table.rows[1].cells[2].text = "{position}{/workHistory}"
                # delete other rows
                for i in range(len(table.rows)-1, 1, -1):
                    table._element.remove(table.rows[i]._element)
        else:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        state = process_p(p, state)
                        
    for i, p in enumerate(doc.paragraphs):
        if "VI. CÁC CÔNG TRÌNH KHOA HỌC" in p.text:
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = "{#scientificWorks}{.}{/scientificWorks}"
        if "VII. KHEN THƯỞNG, KỈ LUẬT" in p.text:
            pass
        if "Khen thưởng: " in p.text:
            doc.paragraphs[i].text = "Khen thưởng: {rewards}"
        if "Kỷ luật: " in p.text:
            doc.paragraphs[i].text = "Kỷ luật: {discipline}"

    doc.save(output_path)
    
convert_user_file('1. Phiếu Lý lịch học viên.docx', 'public/1_Template_Mailing_v4.docx')
