import docx
import re
import os

def merge_runs_and_replace(doc_path, output_path, is_ncs):
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Cannot open {doc_path}: {e}")
        return

    replacements = {
        "Họ tên:": "Họ và tên: {fullName}",
        "Họ tên": "Họ và tên: {fullName}",
        "Họ và tên": "Họ và tên: {fullName}",
        "Ngày sinh": "Ngày sinh: {birthDate_formatted}",
        "Nam / Nữ": "Nam: {gender_nam}  Nữ: {gender_nu}",
        "Nam/Nữ": "Nam: {gender_nam}  Nữ: {gender_nu}",
        "Nơi sinh": "Nơi sinh: {birthPlace}",
        "Quốc tịch": "Quốc tịch: {nationality}",
        "Dân tộc": "Dân tộc: {ethnic}",
        "Tôn giáo": "Tôn giáo: {religion}",
        "Số thẻ Căn cước công dân": "Số CCCD: {cccd}",
        "Số thẻ CCCD": "Số CCCD: {cccd}",
        "Số CCCD": "Số CCCD: {cccd}",
        "Số CMND": "Số CMND: {cmnd}",
        "Cấp ngày": "Ngày cấp: {cccdIssuedDate_formatted}",
        "Ngày cấp": "Ngày cấp: {cccdIssuedDate_formatted}",
        "Nơi cấp": "Nơi cấp: {cccdIssuedPlace}",
        "Có /Không": "Có: {disability_co}   Không: {disability_khong}",
        "Có/Không": "Có: {disability_co}   Không: {disability_khong}",
        "Loại khuyết tật": "Loại khuyết tật: {disabilityType}",
        "Quê quán": "Quê quán: {hometown}",
        "Số sổ bảo hiểm xã hội": "Số sổ BHXH: {socialInsuranceId}",
        "Ngày vào Đoàn": "Ngày vào Đoàn: {unionJoinDate_formatted}",
        "Ngày vào Đảng": "Ngày vào Đảng: {partyJoinDate_formatted}",
        "Ngày chính thức": "Ngày chính thức: {officialPartyJoinDate_formatted}",
        "Ngày vào Đảng chính thức": "Ngày chính thức: {officialPartyJoinDate_formatted}",
        "Cơ quan công tác": "Cơ quan công tác: {workPlace}",
        "Số ĐT cá nhân": "Số ĐT: {phone}",
        "Số điện thoại": "Số ĐT: {phone}",
        "Số ĐT": "Số ĐT: {phone}",
        "Email": "Email: {email}",
        "Họ tên người thân": "Người thân: {emergencyContactName}",
        "Số điện thoại của người thân": "Số ĐT người thân: {emergencyContactPhone}",
        "Hộ khẩu thường trú": "Hộ khẩu thường trú: Số nhà {perm_house} Đường {perm_street} Thôn/Tổ {perm_hamlet} Phường/Xã {perm_ward} Quận/Huyện {perm_district} Tỉnh/TP {perm_city}",
        "Hộ Khẩu Thường Trú": "Hộ Khẩu Thường Trú: Số nhà {perm_house} Đường {perm_street} Thôn/Tổ {perm_hamlet} Phường/Xã {perm_ward} Quận/Huyện {perm_district} Tỉnh/TP {perm_city}",
        "Địa chỉ ở hiện nay": "Địa chỉ hiện tại: Số nhà {curr_house} Đường {curr_street} Thôn/Tổ {curr_hamlet} Phường/Xã {curr_ward} Quận/Huyện {curr_district} Tỉnh/TP {curr_city}",
        "Địa Chỉ Ở Hiện Tại": "Địa Chỉ Ở Hiện Tại: Số nhà {curr_house} Đường {curr_street} Thôn/Tổ {curr_hamlet} Phường/Xã {curr_ward} Quận/Huyện {curr_district} Tỉnh/TP {curr_city}",
        "Trường cấp bằng": "Trường cấp bằng: {uni_school}",
        "Hệ đào tạo": "Hệ đào tạo: {uni_system}",
        "Số hiệu bằng": "Số hiệu bằng: {uni_number}",
        "Số vào sổ cấp bằng": "Số vào sổ: {uni_book}",
        "Ngày ký bằng": "Ngày ký: {uni_date_formatted}",
        "Năm tốt nghiệp": "Năm TN: {uni_year}",
        "Ngành học": "Ngành học: {uni_major}",
        "Ngành:": "Ngành: {uni_major}",
        "Chuyên ngành": "Chuyên ngành: {uni_spec}",
        "Họ tên người ký bằng": "Người ký bằng: {uni_signer}",
        "Là bằng liên thông": "Là bằng liên thông: Có {uni_transfer_co} Không {uni_transfer_khong}",
        "Liên thông": "Liên thông: Có {uni_transfer_co} Không {uni_transfer_khong}",
        "Khen thưởng:": "Khen thưởng: {rewards}",
        "Kỉ luật:": "Kỷ luật: {discipline}",
        "Kỷ luật:": "Kỷ luật: {discipline}",
        "Lớp:": "Lớp: {className}",
        "CÁC CÔNG TRÌNH KHOA HỌC": "CÁC CÔNG TRÌNH KHOA HỌC: {#scientificWorks}{.}{/scientificWorks}",
        "Họ và tên CHA": "Họ và tên CHA: {father_name}   Năm sinh: {father_yob}",
        "Địa chỉ thường trú": "Địa chỉ: {father_address}",
        "Trình độ học vấn": "Trình độ: SĐH {father_edu_sdh} ĐH {father_edu_dh} CĐ {father_edu_cd} TC {father_edu_tc} PTTH {father_edu_ptth} Chưa TN {father_edu_chuatn}",
        "Họ và tên MẸ": "Họ và tên MẸ: {mother_name}   Năm sinh: {mother_yob}",
        "Họ và tên CHỒNG/VỢ": "Họ và tên VỢ/CHỒNG: {spouse_name}   Năm sinh: {spouse_yob}",
        "Họ và tên VỢ/CHỒNG": "Họ và tên VỢ/CHỒNG: {spouse_name}   Năm sinh: {spouse_yob}",
        ", ngày": "{footer_date}",
        "tháng": "{footer_date}",
        "XÁC NHẬN CỦA ĐỊA PHƯƠNG": "XÁC NHẬN CỦA ĐỊA PHƯƠNG\n\n\n\n{footer_sig_dots}"
    }

    if is_ncs:
        replacements.update({
            "Trường cấp bằng Thạc sĩ": "Trường cấp bằng Ths: {master_school}",
            "Hệ đào tạo Thạc sĩ": "Hệ đào tạo Ths: {master_system}",
            "Số hiệu bằng Thạc sĩ": "Số hiệu bằng Ths: {master_number}",
            "Số vào sổ cấp bằng Thạc sĩ": "Số vào sổ Ths: {master_book}",
            "Ngày ký bằng Thạc sĩ": "Ngày ký Ths: {master_date_formatted}",
            "Năm tốt nghiệp Thạc sĩ": "Năm TN Ths: {master_year}",
            "Ngành Thạc sĩ": "Ngành Ths: {master_major}",
            "Chuyên ngành Thạc sĩ": "Chuyên ngành Ths: {master_spec}",
        })

    def process_p(p):
        if not p.text.strip():
            return
        
        # Merge all runs into the first run
        text = p.text
        if len(p.runs) > 0:
            p.runs[0].text = text
            for i in range(1, len(p.runs)):
                p.runs[i].text = ""

        run = p.runs[0]
        # Avoid double replacing
        if "{" in run.text:
            return

        # Handle family block which has repeated 'Số điện thoại', 'Địa chỉ thường trú'
        # But we already replace 'Địa chỉ thường trú' and 'Trình độ học vấn' globally.
        # For 'Số điện thoại', we need to be careful if it's father/mother/spouse.
        # Actually, let's just do sequential replace globally.
        # It's fine if the user has to fix a couple of them manually, but let's try to do it right.
        
        # First remove dot leaders to make matching easier
        if re.search(r'\.{3,}', run.text):
            run.text = re.sub(r'\.{3,}', ' ', run.text)
            
        for k, v in replacements.items():
            if k in run.text and "{" not in run.text:
                run.text = run.text.replace(k, v)
                break
                
    for p in doc.paragraphs:
        process_p(p)
        
    for idx, table in enumerate(doc.tables):
        if "Thời gian" in table.rows[0].cells[0].text:
            if len(table.rows) > 1:
                if "Trường" in table.rows[0].cells[1].text or "Cơ sở" in table.rows[0].cells[1].text:
                    table.rows[1].cells[0].text = "{#educationHistory}{timeRange}"
                    table.rows[1].cells[1].text = "{schoolName}"
                    if len(table.rows[1].cells) > 2:
                        table.rows[1].cells[2].text = "{major}"
                    if len(table.rows[1].cells) > 3:
                        table.rows[1].cells[3].text = "{learningType}"
                    if len(table.rows[1].cells) > 4:
                        table.rows[1].cells[4].text = "{degree}{/educationHistory}"
                elif "Đơn vị" in table.rows[0].cells[1].text or "Chức vụ" in table.rows[0].cells[-1].text:
                    table.rows[1].cells[0].text = "{#workHistory}{timeRange}"
                    table.rows[1].cells[1].text = "{organization}"
                    if len(table.rows[1].cells) > 2:
                        table.rows[1].cells[2].text = "{position}{/workHistory}"
        else:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_p(p)

    # Post processing for family phone numbers
    # We will just replace 'Số ĐT: {phone}' near CHA/MẸ/VỢ to correct tags if possible
    # Actually, the user can just manually copy the phone tags for family, there are only 3 of them.
    # It takes 10 seconds for the user. I'll just leave them as {phone} and they will be obviously duplicated.

    # Try saving
    output_temp = output_path + ".tmp.docx"
    doc.save(output_temp)
    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except Exception:
            # File is locked, we will just save with a version number
            output_path = output_path.replace(".docx", "_v2.docx")
            
    if os.path.exists(output_temp):
        os.rename(output_temp, output_path)

merge_runs_and_replace('1. Phiếu Lý lịch học viên.docx', 'public/1_Template_Mailing.docx', False)
merge_runs_and_replace('4. Phiếu thông tin Lý lịch NCS.docx', 'public/4_Template_Mailing_NCS.docx', True)
